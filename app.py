from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Depends, status, Request, Response
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel, EmailStr
import uvicorn
import os
import json
from datetime import datetime, timedelta
from typing import Optional, List, Dict
from pathlib import Path
import bcrypt
import logging
import uuid
import aiofiles
import shutil
import asyncio
import threading
import time
from transcribe_audio import AudioTranscriber
from docx import Document
import httpx
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load API URLs from environment variables
WHISPER_API_URL = os.getenv("WHISPER_API_URL")
TEXT_API_URL = os.getenv("TEXT_API_URL")

if not WHISPER_API_URL or not TEXT_API_URL:
    raise ValueError("WHISPER_API_URL and TEXT_API_URL must be set in .env file")

# Create necessary directories
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
DATA_DIR = Path("data")
DATA_DIR.mkdir(exist_ok=True)
REPORTS_DIR = Path("reports")
REPORTS_DIR.mkdir(exist_ok=True)

# Data storage files
USERS_FILE = DATA_DIR / "users.json"
UPLOADS_FILE = DATA_DIR / "uploads.json"
TRANSCRIPTS_FILE = DATA_DIR / "transcripts.json"
JOBS_FILE = DATA_DIR / "jobs.json"
REPORTS_FILE = DATA_DIR / "reports.json"

# Initialize storage files if they don't exist
def init_storage_file(file_path: Path, default_data: Dict = None):
    if not file_path.exists():
        with open(file_path, 'w') as f:
            json.dump(default_data or {}, f)

init_storage_file(USERS_FILE)
init_storage_file(UPLOADS_FILE)
init_storage_file(TRANSCRIPTS_FILE)
init_storage_file(JOBS_FILE)
init_storage_file(REPORTS_FILE)

def load_data(file_path: Path) -> Dict:
    try:
        with open(file_path, 'r') as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"Error loading data from {file_path}: {str(e)}")
        return {}

def save_data(file_path: Path, data: Dict):
    try:
        with open(file_path, 'w') as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        logger.error(f"Error saving data to {file_path}: {str(e)}")

# Models
class UserLogin(BaseModel):
    email: str
    password: str

class UserCreate(BaseModel):
    username: str
    email: str
    password: str
    confirm_password: str
    full_name: str

class User(BaseModel):
    username: str
    email: str
    full_name: str

class TranscribeRequest(BaseModel):
    file_id: str
    title: str
    max_workers: int = 6
    model_name: str = "Whisper Malaysia"
    language: str = "auto"

class TranscriptUpdate(BaseModel):
    text: str

class GenerateReportRequest(BaseModel):
    transcript_id: str
    prompt: str
    title: str

# Create FastAPI app instance
app = FastAPI(title="PDRM Meeting Minutes Assistant")

# Add CORS middleware
origins = [
    "http://localhost:3000",
    "http://localhost:3001",
    "http://localhost:5173",
    "http://127.0.0.1:3000",
    "http://127.0.0.1:3001",
    "http://127.0.0.1:5173",
    "http://192.168.50.140:3000",
    "http://192.168.50.140:3001",
    "http://192.168.50.140:5173",
    "https://columns-wide-twice-shape.trycloudflare.com",
    "http://columns-wide-twice-shape.trycloudflare.com",
    "https://*.trycloudflare.com",  # Allow all Cloudflare tunnel domains
    "http://*.trycloudflare.com",   # Allow both http and https for tunnels
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=[
        "Content-Type",
        "Authorization",
        "Accept",
        "Origin",
        "X-Requested-With",
        "Access-Control-Request-Method",
        "Access-Control-Request-Headers"
    ],
)

# Storage dictionaries (now backed by files)
upload_progress = {}  # This can stay in memory as it's temporary

def process_transcription(job_id: str):
    """Process transcription using AudioTranscriber."""
    try:
        jobs = load_data(JOBS_FILE)
        job = jobs[job_id]
        job["status"] = "processing"
        save_data(JOBS_FILE, jobs)
        
        # Create AudioTranscriber instance
        transcriber = AudioTranscriber()
        
        # Get file path and settings
        file_path = job["file_path"]
        settings = job["settings"]
        
        # Update progress for initialization
        job["progress"] = 0
        job["message"] = "Memulakan transkripsi..."
        save_data(JOBS_FILE, jobs)
        
        # Start transcription
        try:
            # Transcribe the file
            transcription = transcriber.transcribe_file(
                audio_path=file_path,
                max_workers=settings["max_workers"],
                model_name=settings["model_name"],
                language=settings["language"],
                format="txt"  # We want plain text output
            )
            
            # Update job status
            job["status"] = "completed"
            job["progress"] = 100
            job["message"] = "Transkrip selesai"
            save_data(JOBS_FILE, jobs)
            
            # Store the transcript
            transcripts = load_data(TRANSCRIPTS_FILE)
            transcripts[job_id] = {
                "text": transcription,
                "title": settings["title"]
            }
            save_data(TRANSCRIPTS_FILE, transcripts)
            
        except Exception as e:
            logger.error(f"Transcription error: {str(e)}")
            job["status"] = "error"
            job["message"] = f"Ralat semasa transkripsi: {str(e)}"
            job["progress"] = 0
            save_data(JOBS_FILE, jobs)
            
    except Exception as e:
        logger.error(f"Error in transcription process: {str(e)}")
        jobs = load_data(JOBS_FILE)
        if job_id in jobs:
            jobs[job_id]["status"] = "error"
            jobs[job_id]["message"] = "Ralat semasa pemprosesan"
            save_data(JOBS_FILE, jobs)

async def generate_report_content(transcript_text: str, prompt: str) -> str:
    """Generate report content using Malaysian text model."""
    try:
        # Log request details
        logger.info(f"Making request to text model at: {TEXT_API_URL}")
        logger.info(f"Transcript length: {len(transcript_text)}")
        logger.info(f"Prompt length: {len(prompt)}")

        # Create client with more specific timeout settings - increased for LLM processing
        timeout = httpx.Timeout(connect=30.0, read=300.0, write=30.0, pool=30.0)
        
        async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
            # Prepare request payload
            payload = {
                "messages": [
                    {
                        "role": "system",
                        "content": """Anda adalah penulis laporan profesional. Tugas anda adalah menganalisis transkrip yang diberikan
                        dan membuat laporan berstruktur mengikut format yang diberikan. Laporan tersebut mestilah jelas,
                        profesional, dan mengikut format yang betul. Gunakan bullet points di mana sesuai."""
                    },
                    {
                        "role": "user",
                        "content": f"Berikut adalah transkrip:\n\n{transcript_text[:4000]}\n\nSila buat laporan mengikut format ini:\n\n{prompt}"
                    }
                ],
                "model": "llm_model",
                "temperature": 0.7,
                "max_tokens": 2000
            }
            
            logger.info(f"Request payload prepared, making POST to: {TEXT_API_URL}/chat/completions")

            # Make request with detailed error handling
            try:
                response = await client.post(
                    f"{TEXT_API_URL}/chat/completions",
                    json=payload,
                    headers={
                        "Content-Type": "application/json",
                        "Accept": "application/json",
                        "User-Agent": "PDRM-Minutes-App/1.0"
                    }
                )
                
                # Log response status
                logger.info(f"Text model response status: {response.status_code}")
                logger.info(f"Response headers: {dict(response.headers)}")
                
                # Handle non-200 responses
                if response.status_code != 200:
                    logger.error(f"Non-200 response: {response.status_code}")
                    logger.error(f"Response content: {response.text}")
                    raise Exception(f"Text model API returned status {response.status_code}")
                
                # Parse response
                try:
                    result = response.json()
                    logger.info("Successfully parsed response JSON")
                    logger.info(f"Response keys: {list(result.keys())}")
                except Exception as e:
                    logger.error(f"Error parsing response JSON: {str(e)}")
                    logger.error(f"Response content: {response.text}")
                    raise Exception("Invalid JSON response from text model")
                
                # Validate response structure
                if "choices" not in result:
                    logger.error(f"Missing 'choices' in response. Available keys: {list(result.keys())}")
                    logger.error(f"Full response: {result}")
                    raise Exception("Invalid response structure from text model")
                
                if not result["choices"]:
                    logger.error("Empty choices array in response")
                    raise Exception("No content generated by text model")
                
                if "message" not in result["choices"][0]:
                    logger.error(f"Missing 'message' in first choice: {result['choices'][0]}")
                    raise Exception("Invalid response format from text model")
                
                if "content" not in result["choices"][0]["message"]:
                    logger.error(f"Missing 'content' in message: {result['choices'][0]['message']}")
                    raise Exception("Invalid response content from text model")
                
                content = result["choices"][0]["message"]["content"]
                logger.info(f"Successfully generated content of length: {len(content)}")
                return content

            except httpx.TimeoutException as e:
                logger.error(f"Timeout error to text model: {str(e)}")
                raise Exception("Request to text model API timed out")
                
            except httpx.HTTPStatusError as e:
                logger.error(f"HTTP error from text model: {e.response.status_code}")
                logger.error(f"Response content: {e.response.text}")
                raise Exception(f"Text model API error: {e.response.status_code}")

            except httpx.ConnectError as e:
                logger.error(f"Connection error to text model: {str(e)}")
                logger.error(f"Full error details: {repr(e)}")
                raise Exception(f"Failed to connect to text model API: {str(e)}")
                
            except httpx.RequestError as e:
                logger.error(f"Request error to text model: {str(e)}")
                logger.error(f"Full error details: {repr(e)}")
                logger.error(f"Error type: {type(e).__name__}")
                raise Exception(f"Request error to text model API: {str(e)}")

    except Exception as e:
        logger.error(f"Error generating report content: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        raise Exception(f"Failed to generate report content: {str(e)}")

def create_docx_report(title: str, prompt: str, content: str) -> Document:
    """Create a DOCX document with the report content."""
    doc = Document()
    
    # Add logo to the top of the document
    try:
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create a paragraph for the logo
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add logo image
        logo_path = Path("static/asset/logo.png")
        if logo_path.exists():
            run = logo_paragraph.runs[0] if logo_paragraph.runs else logo_paragraph.add_run()
            run.add_picture(str(logo_path), width=Inches(2.5))
        else:
            # Fallback if logo not found
            logo_paragraph.add_run("POLIS DIRAJA MALAYSIA").bold = True
            
        # Add some space after logo
        doc.add_paragraph()
        
    except Exception as e:
        logger.warning(f"Could not add logo to document: {str(e)}")
        # Add text header as fallback
        header_para = doc.add_paragraph("POLIS DIRAJA MALAYSIA")
        header_para.runs[0].bold = True
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    # Add title
    title_para = doc.add_heading(f"{title}", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content directly without showing the prompt
    # Split content by lines and add them to document
    for line in content.split('\n'):
        if line.strip():
            # Check if line is a heading (starts with #)
            if line.startswith('#'):
                level = line.count('#')
                text = line.strip('#').strip()
                doc.add_heading(text, level)
            else:
                # Check if line is a bullet point
                if line.strip().startswith('- ') or line.strip().startswith('* '):
                    doc.add_paragraph(line.strip('- ').strip('* '), style='List Bullet')
                else:
                    doc.add_paragraph(line)
    
    return doc

def process_report_generation(report_id: str):
    """Process report generation."""
    try:
        reports = load_data(REPORTS_FILE)
        report = reports[report_id]
        report["status"] = "processing"
        save_data(REPORTS_FILE, reports)

        # Get transcript
        transcripts = load_data(TRANSCRIPTS_FILE)
        transcript = transcripts.get(report["transcript_id"])
        
        if not transcript:
            raise Exception("Transcript not found")

        # Update progress
        reports = load_data(REPORTS_FILE)
        report = reports[report_id]
        report["progress"] = 20
        report["message"] = "Menganalisis transkrip..."
        save_data(REPORTS_FILE, reports)

        # Generate report content using LLM
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        content = loop.run_until_complete(
            generate_report_content(transcript["text"], report["prompt"])
        )
        loop.close()

        # Update progress
        reports = load_data(REPORTS_FILE)
        report = reports[report_id]
        report["progress"] = 60
        report["message"] = "Menjana dokumen laporan..."
        save_data(REPORTS_FILE, reports)

        # Create DOCX document
        doc = create_docx_report(report["title"], report["prompt"], content)
        
        # Save document
        report_path = REPORTS_DIR / f"{report_id}.docx"
        doc.save(str(report_path))

        # Update report status
        reports = load_data(REPORTS_FILE)
        report = reports[report_id]
        report["status"] = "completed"
        report["progress"] = 100
        report["message"] = "Laporan selesai"
        report["file_path"] = str(report_path)
        save_data(REPORTS_FILE, reports)

    except Exception as e:
        logger.error(f"Report generation error: {str(e)}")
        reports = load_data(REPORTS_FILE)
        if report_id in reports:
            reports[report_id]["status"] = "error"
            reports[report_id]["message"] = f"Ralat semasa menjana laporan: {str(e)}"
            reports[report_id]["progress"] = 0
            save_data(REPORTS_FILE, reports)

@app.post("/register", response_model=User)
async def register(user_data: UserCreate):
    try:
        logger.info(f"Received registration request for username: {user_data.username}")
        
        users = load_data(USERS_FILE)
        
        # Validate passwords match
        if user_data.password != user_data.confirm_password:
            logger.warning("Password mismatch")
            raise HTTPException(
                status_code=400,
                detail="Passwords do not match"
            )
        
        # Check if username exists
        if user_data.username in users:
            logger.warning(f"Username {user_data.username} already exists")
            raise HTTPException(
                status_code=400,
                detail="Username already registered"
            )

        # Hash password
        hashed = bcrypt.hashpw(user_data.password.encode(), bcrypt.gensalt())
        
        # Store user
        users[user_data.username] = {
            "username": user_data.username,
            "email": user_data.email,
            "full_name": user_data.full_name,
            "hashed_password": hashed.decode(),  # Convert bytes to string for JSON storage
            "created_at": datetime.now().isoformat(),
            "last_login": None
        }
        save_data(USERS_FILE, users)
        
        logger.info(f"Successfully registered user: {user_data.username}")
        return {
            "username": user_data.username,
            "email": user_data.email,
            "full_name": user_data.full_name
        }
        
    except HTTPException as e:
        raise
        
    except Exception as e:
        logger.error(f"Registration error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="An error occurred during registration"
        )

@app.post("/login")
async def login(user_data: UserLogin):
    logger.info(f"Login attempt for email: {user_data.email}")
    
    users = load_data(USERS_FILE)
    
    # Find user by email
    user = None
    username = None
    for uname, udata in users.items():
        if udata.get("email", "").lower() == user_data.email.lower():
            user = udata
            username = uname
            break
    
    if not user:
        logger.warning(f"User not found with email: {user_data.email}")
        raise HTTPException(
            status_code=401,
            detail="Incorrect email or password"
        )
    
    if not bcrypt.checkpw(user_data.password.encode(), user["hashed_password"].encode()):
        logger.warning(f"Invalid password for email: {user_data.email}")
        raise HTTPException(
            status_code=401,
            detail="Incorrect email or password"
        )
    
    # Update last login
    users[username]["last_login"] = datetime.now().isoformat()
    save_data(USERS_FILE, users)
    
    logger.info(f"Successful login for email: {user_data.email}")
    return {
        "token": "dummy_token",  # In production, use proper JWT
        "name": user["full_name"]
    }

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        # Generate unique file ID
        file_id = str(uuid.uuid4())
        upload_progress[file_id] = 0
        
        # Create file path
        file_path = UPLOAD_DIR / f"{file_id}_{file.filename}"
        logger.info(f"Saving file to: {file_path}")
        
        # Save file with progress tracking
        file_size = 0
        chunk_size = 8192  # 8KB chunks
        
        async with aiofiles.open(file_path, 'wb') as f:
            while chunk := await file.read(chunk_size):
                await f.write(chunk)
                file_size += len(chunk)
                upload_progress[file_id] = min(99, int((file_size / file.size) * 100))
        
        upload_progress[file_id] = 100
        
        # Store file info
        uploads = load_data(UPLOADS_FILE)
        uploads[file_id] = {
            "filename": file.filename,
            "path": str(file_path),
            "size": file_size
        }
        save_data(UPLOADS_FILE, uploads)
        
        logger.info(f"File saved successfully. ID: {file_id}, Path: {file_path}")
        return {
            "file_id": file_id,
            "filename": file.filename,
            "size": file_size
        }
        
    except Exception as e:
        logger.error(f"Upload error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to upload file: {str(e)}"
        )

@app.get("/upload-progress/{file_id}")
async def get_upload_progress(file_id: str):
    if file_id not in upload_progress:
        raise HTTPException(
            status_code=404,
            detail="Upload not found"
        )
    
    return {
        "progress": upload_progress[file_id]
    }

@app.post("/transcribe")
async def transcribe_audio(request: TranscribeRequest):
    try:
        logger.info(f"Received transcription request for file ID: {request.file_id}")
        
        # Check if file exists
        uploads = load_data(UPLOADS_FILE)
        file_info = uploads.get(request.file_id)
        if not file_info:
            logger.error(f"File not found in uploaded_files. Available IDs: {list(uploads.keys())}")
            raise HTTPException(
                status_code=404,
                detail="File not found"
            )

        # Check if file exists on disk
        file_path = Path(file_info["path"])
        if not file_path.exists():
            logger.error(f"File not found on disk: {file_path}")
            raise HTTPException(
                status_code=404,
                detail="File not found on disk"
            )

        # Generate unique request ID
        request_id = str(uuid.uuid4())
        
        # Save transcription job info
        jobs = load_data(JOBS_FILE)
        jobs[request_id] = {
            "status": "pending",
            "file_name": file_info["filename"],
            "file_path": str(file_path),
            "settings": {
                "max_workers": request.max_workers,
                "model_name": request.model_name,
                "language": request.language,
                "title": request.title
            },
            "progress": 0,
            "message": "Memulakan transkripsi..."
        }
        save_data(JOBS_FILE, jobs)
        
        # Start transcription in a background thread
        thread = threading.Thread(target=process_transcription, args=(request_id,))
        thread.daemon = True  # Make thread daemon so it doesn't block shutdown
        thread.start()
        
        logger.info(f"Created transcription job: {request_id}")
        
        return {
            "request_id": request_id,
            "message": "Transcription job started"
        }
        
    except HTTPException as e:
        raise
        
    except Exception as e:
        logger.error(f"Transcription error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to start transcription: {str(e)}"
        )

@app.get("/progress/{request_id}")
async def get_progress(request_id: str):
    jobs = load_data(JOBS_FILE)
    if request_id not in jobs:
        raise HTTPException(
            status_code=404,
            detail="Transcription job not found"
        )
    
    job = jobs[request_id]
    return {
        "status": job["status"],
        "progress": job["progress"],
        "message": job["message"]
    }

@app.get("/transcripts/{transcript_id}")
async def get_transcript(transcript_id: str):
    transcripts = load_data(TRANSCRIPTS_FILE)
    if transcript_id not in transcripts:
        raise HTTPException(
            status_code=404,
            detail="Transcript not found"
        )
    
    return transcripts[transcript_id]

@app.put("/transcripts/{transcript_id}")
async def update_transcript(transcript_id: str, update: TranscriptUpdate):
    transcripts = load_data(TRANSCRIPTS_FILE)
    if transcript_id not in transcripts:
        raise HTTPException(
            status_code=404,
            detail="Transcript not found"
        )
    
    transcripts[transcript_id]["text"] = update.text
    save_data(TRANSCRIPTS_FILE, transcripts)
    return transcripts[transcript_id]

@app.delete("/transcripts/{transcript_id}")
async def delete_transcript(transcript_id: str):
    transcripts = load_data(TRANSCRIPTS_FILE)
    if transcript_id not in transcripts:
        raise HTTPException(
            status_code=404,
            detail="Transcript not found"
        )
    
    # Delete transcript
    del transcripts[transcript_id]
    save_data(TRANSCRIPTS_FILE, transcripts)
    
    # Delete associated job if exists
    jobs = load_data(JOBS_FILE)
    if transcript_id in jobs:
        del jobs[transcript_id]
        save_data(JOBS_FILE, jobs)
    
    # Delete associated file if exists
    uploads = load_data(UPLOADS_FILE)
    if transcript_id in uploads:
        file_info = uploads[transcript_id]
        file_path = Path(file_info["path"])
        if file_path.exists():
            file_path.unlink()
        del uploads[transcript_id]
        save_data(UPLOADS_FILE, uploads)
    
    return {"message": "Transcript deleted successfully"}

@app.post("/generate-report")
async def generate_report(request: GenerateReportRequest):
    try:
        logger.info(f"Received report generation request for transcript ID: {request.transcript_id}")
        
        # Check if transcript exists
        transcripts = load_data(TRANSCRIPTS_FILE)
        if request.transcript_id not in transcripts:
            raise HTTPException(
                status_code=404,
                detail="Transcript not found"
            )

        # Generate unique report ID
        report_id = str(uuid.uuid4())
        
        # Save report job info
        reports = load_data(REPORTS_FILE)
        reports[report_id] = {
            "id": report_id,  # Add ID to the report data
            "status": "pending",
            "transcript_id": request.transcript_id,
            "prompt": request.prompt,
            "title": request.title,
            "progress": 0,
            "message": "Memulakan penjanaan laporan...",
            "created_at": datetime.now().isoformat()
        }
        save_data(REPORTS_FILE, reports)
        
        # Start report generation in a background thread
        thread = threading.Thread(target=process_report_generation, args=(report_id,))
        thread.daemon = True
        thread.start()
        
        logger.info(f"Created report generation job: {report_id}")
        
        return {
            "report_id": report_id,
            "message": "Report generation started"
        }
        
    except HTTPException as e:
        raise
        
    except Exception as e:
        logger.error(f"Report generation error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to start report generation: {str(e)}"
        )

@app.get("/reports/{report_id}/progress")
async def get_report_progress(report_id: str):
    reports = load_data(REPORTS_FILE)
    if report_id not in reports:
        raise HTTPException(
            status_code=404,
            detail="Report not found"
        )
    
    report = reports[report_id]
    return {
        "status": report["status"],
        "progress": report["progress"],
        "message": report["message"]
    }

@app.get("/reports/{report_id}")
async def get_report(report_id: str):
    reports = load_data(REPORTS_FILE)
    if report_id not in reports:
        raise HTTPException(
            status_code=404,
            detail="Report not found"
        )
    
    report = reports[report_id]
    if report["status"] != "completed":
        raise HTTPException(
            status_code=400,
            detail="Report is not ready yet"
        )
    
    file_path = Path(report["file_path"])
    if not file_path.exists():
        raise HTTPException(
            status_code=404,
            detail="Report file not found"
        )
    
    return FileResponse(
        path=file_path,
        filename=f"{report['title']}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.get("/reports")
async def list_reports():
    reports = load_data(REPORTS_FILE)
    return [
        {**report, "id": report_id}  # Add ID to each report
        for report_id, report in reports.items()
    ]

@app.delete("/reports/{report_id}")
async def delete_report(report_id: str):
    reports = load_data(REPORTS_FILE)
    if report_id not in reports:
        raise HTTPException(
            status_code=404,
            detail="Report not found"
        )
    
    report = reports[report_id]
    
    # Delete report file if exists
    if "file_path" in report:
        file_path = Path(report["file_path"])
        if file_path.exists():
            file_path.unlink()
    
    # Delete report record
    del reports[report_id]
    save_data(REPORTS_FILE, reports)
    
    return {"message": "Report deleted successfully"}

@app.post("/logout")
async def logout():
    logger.info("User logged out")
    return {"message": "Successfully logged out"}

@app.get("/statistics")
async def get_statistics():
    """Get comprehensive system statistics."""
    try:
        users = load_data(USERS_FILE)
        uploads = load_data(UPLOADS_FILE)
        transcripts = load_data(TRANSCRIPTS_FILE)
        jobs = load_data(JOBS_FILE)
        reports = load_data(REPORTS_FILE)
        
        # Calculate basic counts
        total_users = len(users)
        total_audio_files = len(uploads)
        total_transcripts = len(transcripts)
        total_reports = len(reports)
        
        # Calculate transcript status breakdown
        transcript_statuses = {}
        for job_id, job in jobs.items():
            status = job.get('status', 'unknown')
            transcript_statuses[status] = transcript_statuses.get(status, 0) + 1
            
        # Calculate report status breakdown
        report_statuses = {}
        for report_id, report in reports.items():
            status = report.get('status', 'unknown')
            report_statuses[status] = report_statuses.get(status, 0) + 1
        
        # Calculate monthly user registrations for the last 12 months
        import calendar
        
        monthly_registrations = {}
        current_date = datetime.now()
        
        for i in range(12):
            # Calculate the month/year for i months ago
            target_date = current_date - timedelta(days=i*30)
            month_key = target_date.strftime("%Y-%m")
            month_name = f"{calendar.month_name[target_date.month]} {target_date.year}"
            monthly_registrations[month_key] = {
                "month": month_name,
                "count": 0,
                "users": []
            }
        
        # Count user registrations by month
        for username, user_data in users.items():
            created_at = user_data.get('created_at', '2024-01-01')
            try:
                user_date = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                month_key = user_date.strftime("%Y-%m")
                if month_key in monthly_registrations:
                    monthly_registrations[month_key]["count"] += 1
                    monthly_registrations[month_key]["users"].append({
                        "username": username,
                        "full_name": user_data.get('full_name', ''),
                        "email": user_data.get('email', ''),
                        "created_at": created_at
                    })
            except:
                pass
        
        # Calculate recent activity (last 30 days)
        recent_uploads = 0
        recent_transcripts = 0
        recent_reports = 0
        
        thirty_days_ago = datetime.now() - timedelta(days=30)
        
        for upload_id, upload in uploads.items():
            try:
                file_path = Path(upload.get('path', ''))
                if file_path.exists():
                    file_time = datetime.fromtimestamp(file_path.stat().st_mtime)
                    if file_time > thirty_days_ago:
                        recent_uploads += 1
            except:
                pass
        
        for job_id, job in jobs.items():
            recent_transcripts += 1 if job.get('status') == 'completed' else 0
            
        for report_id, report in reports.items():
            try:
                created_at = report.get('created_at', '')
                if created_at:
                    report_date = datetime.fromisoformat(created_at)
                    if report_date > thirty_days_ago:
                        recent_reports += 1
            except:
                pass
        
        return {
            "overview": {
                "total_users": total_users,
                "total_audio_files": total_audio_files,
                "total_transcripts": total_transcripts,
                "total_reports": total_reports,
                "recent_uploads": recent_uploads,
                "recent_transcripts": recent_transcripts,
                "recent_reports": recent_reports
            },
            "transcript_statuses": transcript_statuses,
            "report_statuses": report_statuses,
            "monthly_registrations": monthly_registrations,
            "generated_at": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Error getting statistics: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get statistics: {str(e)}"
        )

@app.get("/statistics/users/{period}")
async def get_user_statistics(period: str):
    """Get detailed user statistics for a specific period."""
    try:
        users = load_data(USERS_FILE)
        
        if period == "all":
            filtered_users = users
        else:
            try:
                year_month = period  # Expected format: "2024-11"
                filtered_users = {}
                
                for username, user_data in users.items():
                    created_at = user_data.get('created_at', '2024-01-01')
                    try:
                        user_date = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                        if user_date.strftime("%Y-%m") == year_month:
                            filtered_users[username] = user_data
                    except:
                        pass
            except:
                filtered_users = users
        
        # Calculate statistics for filtered users
        user_stats = []
        for username, user_data in filtered_users.items():
            user_stats.append({
                "username": username,
                "full_name": user_data.get('full_name', ''),
                "email": user_data.get('email', ''),
                "created_at": user_data.get('created_at', ''),
                "last_login": user_data.get('last_login', 'Never')
            })
        
        return {
            "period": period,
            "user_count": len(user_stats),
            "users": user_stats
        }
        
    except Exception as e:
        logger.error(f"Error getting user statistics: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get user statistics: {str(e)}"
        )

# Root endpoint for testing
@app.get("/")
async def root():
    return {"message": "PDRM Meeting Minutes Assistant API"}

if __name__ == "__main__":
    print("\nStarting PDRM Meeting Minutes Assistant server...")
    print("Server will be accessible at:")
    print("- Local: http://localhost:8080")
    print("- Network: http://0.0.0.0:8080")
    print("- API documentation: http://localhost:8080/docs")
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8080,
        reload=True,
        forwarded_allow_ips="*",  # Trust forwarded headers from proxy
        proxy_headers=True        # Enable proxy header support
    )
